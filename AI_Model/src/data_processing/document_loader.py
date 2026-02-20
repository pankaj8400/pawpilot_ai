import os
import json
import logging
from pathlib import Path
from typing import List, Dict, Any, Optional
from datetime import datetime
import uuid
import PyPDF2
import docx
from bs4 import BeautifulSoup

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


class DataIngestionPipeline:
    """
    Pipeline to ingest documents from a folder and extract content for RAG.
    Supports: PDF, DOCX, TXT, JSON, HTML, CSV, MD
    Generates rich metadata for each document.
    """
    
    def __init__(self, input_folder: str, output_folder: str = "ingested_data",
                 default_source: str = "Unknown", default_severity: str = "medium"):
        """
        Initialize the data ingestion pipeline.
        
        Args:
            input_folder: Path to folder containing documents
            output_folder: Path to save processed documents
            default_source: Default source for documents
            default_severity: Default severity level (low/medium/high/critical)
        """
        self.input_folder = Path(input_folder)
        self.output_folder = Path(output_folder)
        self.default_source = default_source
        self.default_severity = default_severity
        self.supported_formats = {
            '.pdf', '.docx', '.txt', '.json', '.html', '.htm', '.csv', '.md'
        }
        
        if not self.input_folder.exists():
            raise ValueError(f"Input folder does not exist: {input_folder}")
        
        self.output_folder.mkdir(parents=True, exist_ok=True)
        logger.info(f"Initialized pipeline - Input: {self.input_folder}, Output: {self.output_folder}")
    
    def get_supported_files(self) -> List[Path]:
        """Get all supported files from input folder."""
        files = []
        for ext in self.supported_formats:
            files.extend(self.input_folder.glob(f"**/*{ext}"))
        
        logger.info(f"Found {len(files)} supported files")
        return sorted(files)
    
    def extract_pdf(self, file_path: Path) -> str:
        """Extract text from PDF file."""
        try:
            text = []
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    if page_text.strip():
                        text.append(page_text)
            
            logger.info(f"Successfully extracted PDF: {file_path.name}")
            return "\n\n".join(text)
        except Exception as e:
            logger.error(f"Error extracting PDF {file_path.name}: {str(e)}")
            return ""
    
    def extract_docx(self, file_path: Path) -> str:
        """Extract text from DOCX file."""
        try:
            doc = docx.Document(file_path)
            text = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text.append(para.text)
            
            for table in doc.tables:
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    text.append(" | ".join(row_data))
            
            logger.info(f"Successfully extracted DOCX: {file_path.name}")
            return "\n".join(text)
        except Exception as e:
            logger.error(f"Error extracting DOCX {file_path.name}: {str(e)}")
            return ""
    
    def extract_txt(self, file_path: Path) -> str:
        """Extract text from TXT file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            logger.info(f"Successfully extracted TXT: {file_path.name}")
            return content
        except Exception as e:
            logger.error(f"Error extracting TXT {file_path.name}: {str(e)}")
            return ""
    
    def extract_json(self, file_path: Path) -> str:
        """Extract text from JSON file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                data = json.load(file)
            
            # Convert JSON to readable text format
            content = json.dumps(data, indent=2)
            logger.info(f"Successfully extracted JSON: {file_path.name}")
            return content
        except Exception as e:
            logger.error(f"Error extracting JSON {file_path.name}: {str(e)}")
            return ""
    
    def extract_html(self, file_path: Path) -> str:
        """Extract text from HTML file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                html_content = file.read()
            
            soup = BeautifulSoup(html_content, 'html.parser')
            for script in soup(['script', 'style']):
                script.decompose()
            
            text = soup.get_text(separator='\n', strip=True)
            logger.info(f"Successfully extracted HTML: {file_path.name}")
            return text
        except Exception as e:
            logger.error(f"Error extracting HTML {file_path.name}: {str(e)}")
            return ""
    
    def extract_csv(self, file_path: Path) -> str:
        """Extract text from CSV file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            logger.info(f"Successfully extracted CSV: {file_path.name}")
            return content
        except Exception as e:
            logger.error(f"Error extracting CSV {file_path.name}: {str(e)}")
            return ""
    
    def extract_markdown(self, file_path: Path) -> str:
        """Extract text from Markdown file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            logger.info(f"Successfully extracted Markdown: {file_path.name}")
            return content
        except Exception as e:
            logger.error(f"Error extracting Markdown {file_path.name}: {str(e)}")
            return ""
    
    def extract_content(self, file_path: Path) -> str:
        """Route file to appropriate extraction method."""
        ext = file_path.suffix.lower()
        
        extraction_methods = {
            '.pdf': self.extract_pdf,
            '.docx': self.extract_docx,
            '.txt': self.extract_txt,
            '.json': self.extract_json,
            '.html': self.extract_html,
            '.htm': self.extract_html,
            '.csv': self.extract_csv,
            '.md': self.extract_markdown,
        }
        
        if ext in extraction_methods:
            return extraction_methods[ext](file_path)
        else:
            logger.warning(f"Unsupported file format: {ext}")
            return ""
    
    def generate_id(self, file_name: str) -> str:
        """Generate unique ID from file name."""
        # Remove extension and replace spaces/special chars with underscores
        base_name = Path(file_name).stem
        sanitized = "".join(c if c.isalnum() else "_" for c in base_name).lower()
        # Append short uuid to ensure uniqueness
        return f"{sanitized}_{uuid.uuid4().hex[:6]}"
    
    def extract_title(self, file_name: str, content: str) -> str:
        """
        Extract or generate title from file name or content.
        
        Args:
            file_name: Name of the file
            content: Extracted content
        
        Returns:
            Title string
        """
        # Try to get title from file name first
        base_name = Path(file_name).stem
        title = base_name.replace("_", " ").replace("-", " ").title()
        
        # If content is available, try to extract first non-empty line as title
        if content:
            lines = [line.strip() for line in content.split('\n') if line.strip()]
            if lines and len(lines[0]) < 100:  # Title should be relatively short
                title = lines[0]
        
        return title
    
    def extract_tags(self, file_name: str, content: str) -> List[str]:
        """
        Generate tags from file name and content.
        
        Args:
            file_name: Name of the file
            content: Extracted content
        
        Returns:
            List of tags
        """
        tags = []
        
        # Extract tags from file name
        base_name = Path(file_name).stem.lower()
        words = [w for w in base_name.split("_") if len(w) > 2]
        tags.extend(words[:3])  # Limit to first 3 words
        
        # Extract common keywords from content (basic approach)
        if content:
            common_keywords = ["overview", "guide", "tutorial", "reference", "document",
                             "report", "analysis", "summary", "introduction", "section"]
            content_lower = content.lower()
            for keyword in common_keywords:
                if keyword in content_lower and keyword not in tags:
                    tags.append(keyword)
                    if len(tags) >= 5:  # Limit total tags
                        break
        
        return list(set(tags))  # Remove duplicates
    
    def truncate_content(self, content: str, max_length: int = 2000) -> str:
        """Truncate content to a reasonable length while preserving meaning."""
        if len(content) <= max_length:
            return content
        
        # Try to truncate at sentence boundary
        truncated = content[:max_length]
        last_period = truncated.rfind('.')
        if last_period > max_length * 0.8:  # If period is reasonably close
            return truncated[:last_period + 1]
        return truncated + "..."
    
    def create_document_object(self, file_path: Path, content: str, 
                              custom_metadata: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
        """
        Create a standardized document object for RAG with rich metadata.
        
        Args:
            file_path: Path to the source file
            content: Extracted content
            custom_metadata: Optional custom metadata fields
        
        Returns:
            Document object with metadata
        """
        full_content = content
        summary_content = self.truncate_content(content, max_length=2000)
        
        doc = {
            "id": self.generate_id(file_path.name),
            "title": self.extract_title(file_path.name, content),
            "content": full_content,
            "summary": summary_content if len(content) > 2000 else content,
            "tags": self.extract_tags(file_path.name, content),
            "source": custom_metadata.get("source", self.default_source) if custom_metadata else self.default_source,
            "date": custom_metadata.get("date", datetime.now().strftime("%Y-%m-%d")) if custom_metadata else datetime.now().strftime("%Y-%m-%d"),
            "severity": custom_metadata.get("severity", self.default_severity) if custom_metadata else self.default_severity,
            "lang": custom_metadata.get("lang", "en") if custom_metadata else "en",
            "file_name": file_path.name,
            "file_type": file_path.suffix.lower(),
            "file_path": str(file_path),
            "content_length": len(content),
            "word_count": len(content.split()),
            "char_count": len(content),
            "ingestion_date": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        }
        
        return doc
    
    def ingest_all_files(self, custom_metadata_map: Optional[Dict[str, Dict[str, Any]]] = None) -> List[Dict[str, Any]]:
        """
        Ingest all supported files from the folder.
        
        Args:
            custom_metadata_map: Optional mapping of file names to custom metadata
                                Example: {"file.txt": {"source": "VCA", "severity": "high"}}
        
        Returns:
            List of document objects
        """
        documents = []
        files = self.get_supported_files()
        
        if not files:
            logger.warning("No supported files found in the input folder")
            return documents
        
        custom_metadata_map = custom_metadata_map or {}
        
        for file_path in files:
            logger.info(f"Processing: {file_path.name}")
            content = self.extract_content(file_path)
            
            if content.strip():
                custom_meta = custom_metadata_map.get(file_path.name)
                doc = self.create_document_object(file_path, content, custom_meta)
                documents.append(doc)
            else:
                logger.warning(f"No content extracted from {file_path.name}")
        
        logger.info(f"Successfully ingested {len(documents)} documents")
        return documents
    
    def save_ingested_data(self, documents: List[Dict[str, Any]], 
                          format: str = "json") -> str:
        """
        Save ingested data to output folder.
        
        Args:
            documents: List of document objects
            format: Output format ('json' or 'jsonl')
        
        Returns:
            Path to saved file
        """
        if format == "json":
            output_file = self.output_folder / "ingested_documents.json"
            with open(output_file, 'w', encoding='utf-8') as f:
                json.dump(documents, f, indent=2, ensure_ascii=False)
        
        elif format == "jsonl":
            output_file = self.output_folder / "ingested_documents.jsonl"
            with open(output_file, 'w', encoding='utf-8') as f:
                for doc in documents:
                    f.write(json.dumps(doc, ensure_ascii=False) + '\n')
        
        logger.info(f"Saved ingested data to {output_file}")
        return str(output_file)
    
    def process(self, save_format: str = "json", 
               custom_metadata_map: Optional[Dict[str, Dict[str, Any]]] = None) -> tuple[List[Dict[str, Any]], str]:
        """
        Run the complete ingestion pipeline.
        
        Args:
            save_format: Output format ('json' or 'jsonl')
            custom_metadata_map: Optional mapping of file names to custom metadata
        
        Returns:
            Tuple of (documents list, output file path)
        """
        logger.info("Starting data ingestion pipeline...")
        documents = self.ingest_all_files(custom_metadata_map)
        output_path = self.save_ingested_data(documents, format=save_format)
        logger.info("Data ingestion pipeline completed successfully!")
        return documents, output_path


# Example usage
if __name__ == "__main__":
    # Define custom metadata for specific files (optional)
    custom_metadata = {
        "hotspots.txt": {
            "source": "VCA / PetMD",
            "severity": "moderate",
            "date": "2025-12-11",
            "lang": "en"
        },
        "another_file.pdf": {
            "source": "Custom Source",
            "severity": "high",
            "date": "2025-12-10",
            "lang": "en"
        }
    }
    
    # Create pipeline instance
    pipeline = DataIngestionPipeline(
        input_folder="./documents",  # Change to your input folder
        output_folder="./ingested_data",
        default_source="Unknown",
        default_severity="medium"
    )
    
    # Run ingestion with custom metadata
    documents, output_file = pipeline.process(
        save_format="json",
        custom_metadata_map=custom_metadata
    )
    
    # Print summary
    print(f"\n{'='*60}")
    print(f"Ingestion Complete!")
    print(f"{'='*60}")
    print(f"Total documents processed: {len(documents)}")
    print(f"Output saved to: {output_file}")
    
    # Print document summaries
    for doc in documents:
        print(f"\n[{doc['id']}]")
        print(f"Title: {doc['title']}")
        print(f"File: {doc['file_name']} ({doc['file_type']})")
        print(f"Tags: {', '.join(doc['tags'])}")
        print(f"Source: {doc['source']} | Severity: {doc['severity']}")
        print(f"Content: {doc['word_count']} words | Length: {doc['content_length']} chars")